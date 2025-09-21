import os
from typing import Dict, Any, List
from PyPDF2 import PdfReader
import pandas as pd
from docx import Document
import io
from logging_config import logger

def _get_extension(filename):
    """
    Return the lowercase extension string of filename, handling if filename is a tuple/list.
    """
    # If filename is tuple/list like (name, ext)
    if isinstance(filename, (tuple, list)) and len(filename) >= 2:
        return str(filename[1]).lower()
    # Otherwise assume it's a string
    return os.path.splitext(str(filename))[1].lower()

class FileProcessor:
    """Handle different file formats for document processing"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._process_txt,
            '.pdf': self._process_pdf,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.docx': self._process_docx,
            '.doc': self._process_docx
        }

    
    def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process uploaded file based on its extension
        Returns: {
            'content': str,
            'metadata': dict,
            'success': bool,
            'error': str (if any)
        }
        """
        logger.info(f"Processing file: {filename}")
        try:
            # Get file extension
            file_ext = _get_extension(filename)

            
            if file_ext not in self.supported_formats:
                error_msg = f"Unsupported file format: {file_ext}. Supported formats: {list(self.supported_formats.keys())}"
                logger.error(error_msg)
                return {
                    'content': '',
                    'metadata': {},
                    'success': False,
                    'error': error_msg
                }
            
            # Process file based on extension
            processor = self.supported_formats[file_ext]
            logger.info(f"Using processor: {processor.__name__}")
            content = processor(file_content, filename)
            
            logger.info(f"File processed successfully: {filename}")
            return {
                'content': content,
                'metadata': {
                    'filename': filename,
                    'file_type': file_ext,
                    'file_size': len(file_content),
                    'content_length': len(content)
                },
                'success': True,
                'error': None
            }
            
        except Exception as e:
            error_msg = f"Error processing file {filename}: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return {
                'content': '',
                'metadata': {},
                'success': False,
                'error': error_msg
            }
    
    def _process_txt(self, file_content: bytes, filename: str) -> str:
        """Process plain text files"""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decoding failed for {filename}. Trying other encodings.")
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            logger.error(f"Unable to decode text file with common encodings: {filename}")
            raise Exception("Unable to decode text file with common encodings")
    
    def _process_pdf(self, file_content: bytes, filename: str) -> str:
        """Process PDF files using PyPDF2"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---")
                    text_content.append(page_text)
            
            if not text_content:
                logger.warning(f"No text content found in PDF: {filename}")
                raise Exception("No text content found in PDF")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {str(e)}", exc_info=True)
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def _process_excel(self, file_content: bytes, filename: str) -> str:
        """Process Excel files (.xlsx, .xls)"""
        try:
            excel_file = io.BytesIO(file_content)
            
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text_content = []
            for sheet_name, df in excel_data.items():
                text_content.append(f"--- Sheet: {sheet_name} ---")
                
                # Convert DataFrame to readable text
                if not df.empty:
                    # Include column headers
                    headers = "Columns: " + ", ".join(df.columns.astype(str))
                    text_content.append(headers)
                    
                    # Convert rows to text (limit to first 1000 rows to avoid huge content)
                    max_rows = min(1000, len(df))
                    for index, row in df.head(max_rows).iterrows():
                        row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                        if row_text.strip():
                            text_content.append(f"Row {index + 1}: {row_text}")
                    
                    if len(df) > max_rows:
                        text_content.append(f"... (truncated, showing first {max_rows} of {len(df)} rows)")
                else:
                    text_content.append("(Empty sheet)")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing Excel file {filename}: {str(e)}", exc_info=True)
            raise Exception(f"Error processing Excel file: {str(e)}")
    
    def _process_docx(self, file_content: bytes, filename: str) -> str:
        """Process Word documents (.docx)"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("--- Table ---")
                    text_content.extend(table_text)
            
            if not text_content:
                logger.warning(f"No text content found in Word document: {filename}")
                raise Exception("No text content found in Word document")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing Word document {filename}: {str(e)}", exc_info=True)
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def get_supported_formats(self) -> List[str]:
        """Return list of supported file formats"""
        return list(self.supported_formats.keys())

# Initialize file processor
file_processor = FileProcessor()